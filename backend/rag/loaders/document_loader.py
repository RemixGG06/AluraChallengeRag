"""Carga de documentos multiformato para MentorTI Nexus.

Formatos soportados: PDF, DOCX, XLSX, CSV, MD, HTML, TXT.

Decisión de diseño: se usan extractores propios en lugar de los loaders de
langchain-community para (1) no agregar dependencias extra, (2) tener control
total sobre la metadata corporativa (departamento, página, extensión) y
(3) evitar la API deprecada de langchain-community (KB-01).
"""

from __future__ import annotations

import csv
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openpyxl import load_workbook
from pypdf import PdfReader

from backend.config.settings import settings

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".csv", ".md", ".html", ".htm", ".txt"}


# ---------------------------------------------------------------------------
# Extractores de texto por formato. Cada uno retorna lista de tuplas
# (texto, metadata_extra).
# ---------------------------------------------------------------------------

def _extract_pdf(path: Path) -> list[tuple[str, dict]]:
    reader = PdfReader(str(path))
    results = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if text:
            results.append((text, {"page": i + 1}))
    return results


def _extract_docx(path: Path) -> list[tuple[str, dict]]:
    doc = DocxDocument(str(path))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [(text, {})] if text.strip() else []


def _extract_xlsx(path: Path) -> list[tuple[str, dict]]:
    workbook = load_workbook(str(path), read_only=True, data_only=True)
    sheets = []
    for worksheet in workbook.worksheets:
        rows = []
        for row in worksheet.iter_rows(values_only=True):
            values = [str(cell) for cell in row if cell is not None]
            if values:
                rows.append(" | ".join(values))
        if rows:
            sheets.append(f"Hoja: {worksheet.title}\n" + "\n".join(rows))
    workbook.close()
    return [("\n\n".join(sheets), {})] if sheets else []


def _extract_csv(path: Path) -> list[tuple[str, dict]]:
    with open(path, newline="", encoding="utf-8-sig") as file:
        rows = [
            " | ".join(row)
            for row in csv.reader(file)
            if any(cell.strip() for cell in row)
        ]
    return [("\n".join(rows), {})] if rows else []


def _extract_html(path: Path) -> list[tuple[str, dict]]:
    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "lxml")
    for tag in soup(["script", "style"]):
        tag.decompose()
    lines = [line.strip() for line in soup.get_text(separator="\n").splitlines() if line.strip()]
    return [("\n".join(lines), {})] if lines else []


def _extract_text(path: Path) -> list[tuple[str, dict]]:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    return [(text, {})] if text else []


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".xlsx": _extract_xlsx,
    ".csv": _extract_csv,
    ".md": _extract_text,
    ".txt": _extract_text,
    ".html": _extract_html,
    ".htm": _extract_html,
}


# ---------------------------------------------------------------------------
# API pública
# ---------------------------------------------------------------------------

def load_document(file_path: str | Path, department: str = "general") -> list[Document]:
    """Carga un archivo y retorna documentos LangChain con metadata corporativa.

    Metadata incluida: source (nombre de archivo), extension, department y
    page (solo PDF).

    Raises:
        ValueError: si la extensión no está soportada.
        FileNotFoundError: si el archivo no existe.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Archivo no encontrado: {path}")
    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Formato no soportado: '{extension}'. "
            f"Soportados: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    base_metadata = {
        "source": path.name,
        "extension": extension,
        "department": department,
    }
    return [
        Document(page_content=text, metadata={**base_metadata, **extra})
        for text, extra in _EXTRACTORS[extension](path)
    ]


def load_directory(
    directory: str | Path, department: str = "general"
) -> tuple[list[Document], list[str]]:
    """Carga todos los archivos soportados de un directorio.

    Retorna (documentos, errores): los archivos que fallen se reportan como
    texto de error sin detener el proceso.
    """
    directory = Path(directory)
    documents: list[Document] = []
    errors: list[str] = []
    for path in sorted(directory.iterdir()):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            documents.extend(load_document(path, department=department))
        except Exception as exc:  # noqa: BLE001 - se reporta y se continúa
            errors.append(f"{path.name}: {exc}")
    return documents, errors


def split_documents(
    documents: list[Document],
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[Document]:
    """Divide documentos en chunks preservando su metadata."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size or settings.chunk_size,
        chunk_overlap=chunk_overlap or settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(documents)
