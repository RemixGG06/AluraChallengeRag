"""Tests del cargador de documentos multiformato."""

import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook
from pypdf import PdfWriter

from backend.rag.loaders.document_loader import (
    SUPPORTED_EXTENSIONS,
    load_directory,
    load_document,
    split_documents,
)


@pytest.fixture()
def md_file(tmp_path):
    path = tmp_path / "manual_red.md"
    path.write_text(
        "# Manual de Red\n\nLa VLAN de invitados es la 30.\n\n" * 30,
        encoding="utf-8",
    )
    return path


@pytest.fixture()
def csv_file(tmp_path):
    path = tmp_path / "inventario.csv"
    path.write_text(
        "equipo,ip,ubicacion\nrouter-01,10.0.0.1,Data Center\npc-22,10.0.0.22,Piso 3\n",
        encoding="utf-8",
    )
    return path


@pytest.fixture()
def html_file(tmp_path):
    path = tmp_path / "normativa.html"
    path.write_text(
        "<html><head><style>body{color:red}</style></head>"
        "<body><h1>Normativa de Seguridad</h1>"
        "<p>Toda contraseña debe rotarse cada 90 días.</p>"
        "<script>alert('x')</script></body></html>",
        encoding="utf-8",
    )
    return path


@pytest.fixture()
def docx_file(tmp_path):
    path = tmp_path / "onboarding.docx"
    doc = DocxDocument()
    doc.add_paragraph("Guía de onboarding de practicantes de TI.")
    doc.add_paragraph("Solicita tus credenciales al Jefe de TI el primer día.")
    doc.save(path)
    return path


@pytest.fixture()
def xlsx_file(tmp_path):
    path = tmp_path / "contactos.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sistemas"
    sheet.append(["nombre", "extension"])
    sheet.append(["Mesa de ayuda", "4357"])
    workbook.save(path)
    return path


@pytest.fixture()
def pdf_file(tmp_path):
    # PDF en blanco: valida el enrutamiento del extractor sin errores
    path = tmp_path / "vacio.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with open(path, "wb") as file:
        writer.write(file)
    return path


class TestLoadDocument:
    def test_markdown_con_metadata(self, md_file):
        docs = load_document(md_file, department="operaciones")
        assert len(docs) == 1
        assert "Manual de Red" in docs[0].page_content
        assert docs[0].metadata["source"] == "manual_red.md"
        assert docs[0].metadata["extension"] == ".md"
        assert docs[0].metadata["department"] == "operaciones"

    def test_csv(self, csv_file):
        docs = load_document(csv_file, department="sistemas")
        assert len(docs) == 1
        assert "router-01" in docs[0].page_content
        assert "10.0.0.1" in docs[0].page_content

    def test_html_sin_scripts_ni_estilos(self, html_file):
        docs = load_document(html_file, department="rh")
        assert len(docs) == 1
        assert "rotarse cada 90 días" in docs[0].page_content
        assert "alert" not in docs[0].page_content
        assert "color:red" not in docs[0].page_content

    def test_docx(self, docx_file):
        docs = load_document(docx_file)
        assert len(docs) == 1
        assert "onboarding" in docs[0].page_content

    def test_xlsx(self, xlsx_file):
        docs = load_document(xlsx_file)
        assert len(docs) == 1
        assert "Mesa de ayuda" in docs[0].page_content
        assert "Hoja: Sistemas" in docs[0].page_content

    def test_pdf_en_blanco_no_falla(self, pdf_file):
        docs = load_document(pdf_file)
        assert docs == []

    def test_extension_no_soportada(self, tmp_path):
        path = tmp_path / "archivo.exe"
        path.write_bytes(b"MZ")
        with pytest.raises(ValueError, match="Formato no soportado"):
            load_document(path)

    def test_archivo_inexistente(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            load_document(tmp_path / "no_existe.pdf")


class TestLoadDirectory:
    def test_carga_multiple_y_reporta_errores(self, tmp_path, md_file, csv_file):
        # md_file y csv_file viven en tmp_path; agregamos un archivo roto
        (tmp_path / "roto.xlsx").write_bytes(b"esto no es un excel")
        (tmp_path / "ignorado.exe").write_bytes(b"MZ")

        docs, errors = load_directory(tmp_path, department="sistemas")
        sources = {doc.metadata["source"] for doc in docs}

        assert "manual_red.md" in sources
        assert "inventario.csv" in sources
        assert len(errors) == 1
        assert "roto.xlsx" in errors[0]
        assert all(doc.metadata["department"] == "sistemas" for doc in docs)


class TestSplitDocuments:
    def test_genera_chunks_y_conserva_metadata(self, md_file):
        docs = load_document(md_file, department="operaciones")
        chunks = split_documents(docs, chunk_size=200, chunk_overlap=40)
        assert len(chunks) > 1
        assert all(len(chunk.page_content) <= 200 for chunk in chunks)
        assert all(
            chunk.metadata["source"] == "manual_red.md" for chunk in chunks
        )

    def test_documento_corto_un_solo_chunk(self, csv_file):
        docs = load_document(csv_file)
        chunks = split_documents(docs)
        assert len(chunks) == 1


def test_formatos_soportados_del_challenge():
    # Cobertura de dominios del reto Alura: PDF, Word, Excel, Markdown, CSV, HTML
    assert {".pdf", ".docx", ".xlsx", ".md", ".csv", ".html"} <= SUPPORTED_EXTENSIONS
