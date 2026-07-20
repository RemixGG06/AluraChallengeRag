"""Genera los documentos de prueba simulados de MentorTI Nexus.

Cobertura de dominios del challenge Alura:
- RH:         normativa de seguridad TI (.docx) + guía de onboarding (.pdf)
- Operaciones: manual de VPN/red (.md) + procedimiento de backups (.html)
- Sistemas:    inventario de equipos (.xlsx) + contactos de soporte (.csv)

Uso:
    python -m backend.scripts.generate_samples

Nota: el PDF requiere reportlab (herramienta de desarrollo, NO va en
requirements.txt):  pip install reportlab
"""

from __future__ import annotations

import csv
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

from docx import Document as DocxDocument
from openpyxl import Workbook

SAMPLES_DIR = Path(__file__).resolve().parent.parent.parent / "data" / "samples"


def _manual_vpn(path: Path) -> None:
    path.write_text(
        """# Manual de Red Corporativa — Departamento de Operaciones

## 1. Conexión a la VPN

Para conectarte a la VPN corporativa sigue estos pasos:

1. Instala el cliente **GlobalProtect** desde el portal de software interno.
2. Abre el cliente e ingresa el portal: `vpn.empresa.com`.
3. Usa tus credenciales de Active Directory (el mismo usuario y contraseña del
   correo corporativo).
4. Si la conexión falla con error de autenticación, verifica que tu contraseña
   no esté vencida y reinicia el servicio `PanGPS` en tu equipo.
5. Si el problema persiste más de 15 minutos, escala al NOC (extensión 4321).

## 2. WiFi corporativo

- Red **EMPRESA-CORP**: solo para equipos de dominio. Se conecta
  automáticamente con certificado.
- Red **EMPRESA-GUEST**: para visitas y dispositivos personales. La contraseña
  se rota **cada lunes a las 08:00** y se publica en la cartelera digital del
  piso 2 y en el canal #ti-anuncios de Slack.

## 3. VLANs de la red

| VLAN | Uso | Segmento |
|------|-----|----------|
| 10 | Usuarios corporativos | 10.10.10.0/24 |
| 20 | Servidores de producción | 10.10.20.0/24 |
| 30 | Invitados | 10.10.30.0/24 |
| 40 | Voz IP | 10.10.40.0/24 |

## 4. Impresoras de red

Las impresoras siguen la nomenclatura `IMP-PISO-NUMERO`. Ejemplo:
`IMP-P3-02` es la impresora 2 del piso 3. Se instalan desde
`\\\\printsrv\\impresoras` con doble clic.
""",
        encoding="utf-8",
    )


def _normativa_seguridad(path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("Normativa de Seguridad TI — Recursos Humanos", level=1)
    doc.add_paragraph(
        "Documento normativo aplicable a todos los colaboradores, practicantes "
        "y personal tercerizado con acceso a sistemas de la empresa."
    )
    doc.add_heading("1. Política de contraseñas", level=2)
    for item in [
        "Longitud mínima: 12 caracteres, incluyendo mayúsculas, minúsculas, "
        "números y un símbolo.",
        "Rotación obligatoria cada 90 días. El sistema bloquea las últimas 5 "
        "contraseñas usadas.",
        "Está prohibido compartir contraseñas, incluso con el Jefe de TI o "
        "soporte técnico. Nadie del área de TI te pedirá jamás tu contraseña.",
        "Tras 5 intentos fallidos de inicio de sesión, la cuenta se bloquea "
        "por 30 minutos.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("2. Credenciales de primer ingreso", level=2)
    doc.add_paragraph(
        "El primer día, el practicante recibe sus credenciales temporales en "
        "sobre cerrado entregado por RRHH. Debe cambiar la contraseña temporal "
        "en el primer inicio de sesión. Las credenciales de Active Directory "
        "habilitan correo, VPN y portal de RRHH."
    )
    doc.add_heading("3. Uso aceptable de equipos", level=2)
    doc.add_paragraph(
        "Los equipos corporativos son herramientas de trabajo. No está "
        "permitido instalar software sin autorización del área de Sistemas ni "
        "conectar dispositivos USB personales. Todo equipo se entrega con "
        "cifrado de disco activado (BitLocker)."
    )
    doc.add_heading("4. Reporte de incidentes", level=2)
    doc.add_paragraph(
        "Ante cualquier sospecha de phishing, pérdida de equipo o acceso no "
        "autorizado, el colaborador debe reportar dentro de las primeras 2 "
        "horas al correo seguridad@empresa.com o a la extensión 4357."
    )
    doc.save(path)


def _inventario_equipos(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Servidores"
    sheet.append(["hostname", "ip", "rol", "ubicacion", "so"])
    for row in [
        ("srv-dc01", "10.10.20.10", "Controlador de dominio / AD", "Data Center Rack A", "Windows Server 2022"),
        ("srv-files01", "10.10.20.15", "Servidor de archivos corporativo", "Data Center Rack A", "Windows Server 2022"),
        ("srv-backup01", "10.10.20.20", "Servidor de backups Veeam", "Data Center Rack B", "Windows Server 2022"),
        ("srv-print01", "10.10.20.25", "Servidor de impresión", "Data Center Rack B", "Windows Server 2019"),
        ("fw-edge01", "10.10.20.1", "Firewall perimetral / VPN", "Data Center Rack A", "FortiOS"),
    ]:
        sheet.append(row)
    sheet2 = workbook.create_sheet("Equipos practicantes")
    sheet2.append(["etiqueta", "modelo", "usuario_asignado", "estado"])
    for row in [
        ("LPT-TI-011", "Dell Latitude 5420", "Por asignar", "Disponible"),
        ("LPT-TI-012", "Dell Latitude 5420", "Por asignar", "Disponible"),
        ("LPT-TI-013", "HP ProBook 450", "Practicante TI-2026-01", "Asignado"),
        ("MON-TI-030", "Dell P2422H", "Almacén TI", "Disponible"),
    ]:
        sheet2.append(row)
    workbook.save(path)


def _contactos_soporte(path: Path) -> None:
    with open(path, "w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(["area", "contacto", "extension", "correo", "horario"])
        writer.writerows([
            ("Mesa de ayuda TI", "Mesa de Ayuda Nivel 1", "4357", "helpdesk@empresa.com", "Lun-Vie 8:00-18:00"),
            ("NOC / Redes", "Centro de Operaciones de Red", "4321", "noc@empresa.com", "24/7"),
            ("Seguridad informática", "Oficial de Seguridad", "4400", "seguridad@empresa.com", "Lun-Vie 9:00-17:00"),
            ("Sistemas / Inventario", "Administrador de Sistemas", "4415", "sysadmin@empresa.com", "Lun-Vie 8:00-17:00"),
            ("RRHH", "Analista de Personas", "4200", "rrhh@empresa.com", "Lun-Vie 9:00-17:00"),
        ])


def _procedimiento_backups(path: Path) -> None:
    path.write_text(
        """<!DOCTYPE html>
<html lang="es">
<head><meta charset="utf-8"><title>Procedimiento de Backups</title></head>
<body>
<h1>Procedimiento de Backups — Operaciones</h1>
<h2>Política general</h2>
<p>Los backups se rigen por la regla <strong>3-2-1</strong>: 3 copias de los
datos, en 2 medios distintos, con 1 copia fuera del sitio (nube).</p>
<h2>Programación</h2>
<ul>
  <li><strong>Incremental:</strong> lunes a sábado a las 21:00.</li>
  <li><strong>Completo:</strong> domingos a las 02:00.</li>
  <li><strong>Retención:</strong> 14 restore points diarios, 4 semanales,
  12 mensuales.</li>
</ul>
<h2>Restauración de un archivo</h2>
<ol>
  <li>Abrir la consola Veeam en <code>srv-backup01</code> (10.10.20.20).</li>
  <li>Seleccionar <em>Restore &gt; File Level Restore</em>.</li>
  <li>Elegir el restore point y el servidor origen (habitualmente
  <code>srv-files01</code>).</li>
  <li>Copiar el archivo restaurado a la ruta original o a
  <code>RESTAURADOS</code> en la raíz del recurso compartido.</li>
  <li>Notificar al usuario y registrar el ticket.</li>
</ol>
<h2>Verificación diaria</h2>
<p>Cada mañana el analista de turno revisa el reporte de jobs en la consola.
Un job en estado <em>Warning</em> dos días seguidos se escala al
Administrador de Sistemas (extensión 4415).</p>
</body>
</html>
""",
        encoding="utf-8",
    )


def _guia_onboarding(path: Path) -> None:
    """Genera la guía de onboarding en PDF (requiere reportlab, dev-only)."""
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.units import cm
        from reportlab.pdfgen import canvas
    except ImportError:
        print("  ! reportlab no instalado: se omite el PDF "
              "(pip install reportlab para generarlo)")
        return

    pdf = canvas.Canvas(str(path), pagesize=LETTER)
    width, height = LETTER
    y = height - 2 * cm
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(2 * cm, y, "Guia de Onboarding de Practicantes TI - RRHH")
    y -= 1.5 * cm
    pdf.setFont("Helvetica", 11)
    lineas = [
        "Bienvenido al area de Tecnologias de Informacion.",
        "",
        "Primer dia:",
        "1. Recoge tus credenciales temporales en RRHH (piso 1, sobre cerrado).",
        "2. Retira tu laptop asignada en el almacen de Sistemas (etiqueta LPT-TI).",
        "3. Cambia la contrasena temporal en tu primer inicio de sesion.",
        "4. Instala GlobalProtect y valida tu acceso a la VPN (vpn.empresa.com).",
        "5. Unete al canal #ti-anuncios y #practicantes en Slack.",
        "",
        "Primera semana:",
        "- Completa el curso de seguridad informatica en el portal de RRHH.",
        "- Lee el Manual de Red y la Normativa de Seguridad TI.",
        "- Agenda una reunion de induccion con tu mentor asignado.",
        "- Solicita acceso a los sistemas que necesites via ticket en",
        "  helpdesk@empresa.com (extension 4357).",
        "",
        "Contactos clave:",
        "- Mesa de ayuda: extension 4357 (Lun-Vie 8:00-18:00).",
        "- NOC / Redes: extension 4321 (24/7).",
        "- RRHH: extension 4200.",
    ]
    for linea in lineas:
        pdf.drawString(2 * cm, y, linea)
        y -= 0.7 * cm
    pdf.save()


def main() -> None:
    SAMPLES_DIR.mkdir(parents=True, exist_ok=True)
    generadores = {
        "manual_vpn_red.md": _manual_vpn,
        "normativa_seguridad_ti.docx": _normativa_seguridad,
        "inventario_equipos.xlsx": _inventario_equipos,
        "contactos_soporte.csv": _contactos_soporte,
        "procedimiento_backups.html": _procedimiento_backups,
        "guia_onboarding.pdf": _guia_onboarding,
    }
    for nombre, generador in generadores.items():
        destino = SAMPLES_DIR / nombre
        generador(destino)
        estado = "OK" if destino.exists() else "OMITIDO"
        print(f"  [{estado}] {nombre}")
    print(f"\nDocumentos en: {SAMPLES_DIR}")


if __name__ == "__main__":
    main()
